
import streamlit as st
from collections import Counter
import docx
import re
from io import BytesIO
from docx import Document

st.title("تحليل الكلمات القانونية الإدارية في ملفات Word")

uploaded_files = st.file_uploader("ارفع ملفات Word (DOCX)", type=["docx"], accept_multiple_files=True)

# قائمة الكلمات القانونية في القضاء الإداري فقط
admin_keywords = ['قرار', 'إلغاء', 'تعيين', 'ترقية', 'موظف', 'وظيفة', 'وزارة', 'جهة', 'جهة إدارية', 'جهة حكومية', 'مجلس', 'لجنة', 'صلاحية', 'اختصاص', 'وظيفي', 'خدمة', 'وظيفة عامة', 'تعسف', 'سلطة تقديرية', 'دعوى إدارية', 'نزاع إداري', 'لائحة', 'قانون الخدمة', 'مخالفة إدارية', 'جزاء', 'توقيف', 'إعفاء', 'استقالة', 'إحالة', 'تقاعد', 'تعويض', 'مسؤولية إدارية', 'توجيه', 'خطاب', 'عزل', 'إنهاء', 'توصية', 'تكليف', 'تفويض', 'ترخيص', 'إذن', 'نظام', 'إدارة', 'هيئة', 'مؤسسة', 'انتداب', 'ندب', 'عقد إداري', 'مناقصة', 'مناقصات', 'مناقصة عامة', 'استبعاد', 'لجنة مناقصات', 'جهة رقابية', 'رقابة', 'تفتيش', 'رقابة إدارية', 'ديوان المظالم', 'محكمة إدارية', 'هيئة قضاء إداري', 'قرار نهائي', 'قرار ابتدائي']

if uploaded_files:
    all_results = []
    for uploaded_file in uploaded_files:
        doc = docx.Document(uploaded_file)
        full_text = " ".join([para.text for para in doc.paragraphs])
        cleaned_text = re.sub(r'[^؀-ۿ\s]', '', full_text)
        words = cleaned_text.split()
        legal_words = [w for w in words if w in admin_keywords]
        word_counts = Counter(legal_words)
        top_words = word_counts.most_common(10)

        st.write(f"**أكثر الكلمات القانونية الإدارية تكرارًا في الملف {uploaded_file.name}:**")
        for word, count in top_words:
            st.write(f"- {word}: {count} مرة")

        result_text = f"نتائج الملف: {uploaded_file.name}\n"
        for word, count in top_words:
            result_text += f"- {word}: {count} مرة\n"
        result_text += "\n"
        all_results.append(result_text)

    if st.button("تحميل النتائج في ملف Word"):
        export_doc = Document()
        export_doc.add_heading("نتائج الكلمات القانونية الإدارية", level=1)
        for result in all_results:
            export_doc.add_paragraph(result)

        buffer = BytesIO()
        export_doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="اضغط هنا لتحميل النتائج",
            data=buffer,
            file_name="تحليل_القضاء_الاداري.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
